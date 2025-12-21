"""Document parsing service for PDF and DOCX files.

Extracts text content from uploaded files for contract analysis.
"""

import base64
import binascii
import io
from typing import Any

import structlog
from docx import Document as DocxDocument
from pypdf import PdfReader

logger = structlog.get_logger()


class DocumentParserError(Exception):
    """Raised when document parsing fails."""

    pass


class UnsupportedFileTypeError(DocumentParserError):
    """Raised when file type is not supported."""

    pass


# Supported MIME types and their file types
SUPPORTED_MIME_TYPES: dict[str, str] = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/msword": "doc",  # Legacy .doc format (limited support)
    "text/plain": "txt",
}


async def parse_document(file_part: dict[str, Any]) -> str:
    """Parse a document file and extract text content.

    Args:
        file_part: FilePart.file dictionary containing:
            - uri: Data URI (data:<mime>;base64,<data>) or external URL
            - mimeType: MIME type of the file
            - name: Optional filename

    Returns:
        Extracted text content from the document

    Raises:
        DocumentParserError: If parsing fails
        UnsupportedFileTypeError: If file type is not supported
    """
    mime_type = file_part.get("mimeType", "")
    uri = file_part.get("uri", "")
    name = file_part.get("name", "unknown")

    logger.info(
        "Parsing document",
        filename=name,
        mime_type=mime_type,
    )

    # Validate MIME type
    if mime_type not in SUPPORTED_MIME_TYPES:
        raise UnsupportedFileTypeError(
            f"Unsupported file type: {mime_type}. "
            f"Supported types: {', '.join(SUPPORTED_MIME_TYPES.keys())}"
        )

    # Extract file bytes from data URI
    if not uri.startswith("data:"):
        raise DocumentParserError(
            "Only data URIs are currently supported. "
            "External URLs require additional implementation."
        )

    try:
        # Parse data URI: data:<mime>;base64,<data>
        header, encoded_data = uri.split(",", 1)
        file_bytes = base64.b64decode(encoded_data)
    except ValueError as e:
        raise DocumentParserError(f"Invalid data URI format: {e}")
    except binascii.Error as e:
        raise DocumentParserError(f"Failed to decode base64 data: {e}")

    # Parse based on file type
    file_type = SUPPORTED_MIME_TYPES[mime_type]

    if file_type == "pdf":
        return await _parse_pdf(file_bytes, name)
    elif file_type in ("docx", "doc"):
        return await _parse_docx(file_bytes, name)
    elif file_type == "txt":
        return await _parse_text(file_bytes, name)
    else:
        raise UnsupportedFileTypeError(f"Parser not implemented for: {file_type}")


async def _parse_pdf(file_bytes: bytes, filename: str) -> str:
    """Extract text from PDF file.

    Args:
        file_bytes: Raw PDF file bytes
        filename: Original filename for logging

    Returns:
        Extracted text from all pages
    """
    try:
        pdf_file = io.BytesIO(file_bytes)
        reader = PdfReader(pdf_file)

        text_parts: list[str] = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(f"--- Page {i + 1} ---\n{page_text}")

        full_text = "\n\n".join(text_parts)

        logger.info(
            "PDF parsed successfully",
            filename=filename,
            pages=len(reader.pages),
            chars=len(full_text),
        )

        return full_text

    except Exception as e:
        logger.error("PDF parsing failed", filename=filename, error=str(e))
        raise DocumentParserError(f"Failed to parse PDF '{filename}': {e}")


async def _parse_docx(file_bytes: bytes, filename: str) -> str:
    """Extract text from DOCX file.

    Args:
        file_bytes: Raw DOCX file bytes
        filename: Original filename for logging

    Returns:
        Extracted text from all paragraphs and tables
    """
    try:
        docx_file = io.BytesIO(file_bytes)
        doc = DocxDocument(docx_file)

        text_parts: list[str] = []

        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)

        full_text = "\n\n".join(text_parts)

        logger.info(
            "DOCX parsed successfully",
            filename=filename,
            paragraphs=len(doc.paragraphs),
            tables=len(doc.tables),
            chars=len(full_text),
        )

        return full_text

    except Exception as e:
        logger.error("DOCX parsing failed", filename=filename, error=str(e))
        raise DocumentParserError(f"Failed to parse DOCX '{filename}': {e}")


async def _parse_text(file_bytes: bytes, filename: str) -> str:
    """Extract text from plain text file.

    Args:
        file_bytes: Raw file bytes
        filename: Original filename for logging

    Returns:
        Decoded text content
    """
    # Try common encodings in order of likelihood
    encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]

    for encoding in encodings:
        try:
            text = file_bytes.decode(encoding)
            logger.info(
                "Text file parsed",
                filename=filename,
                encoding=encoding,
                chars=len(text),
            )
            return text
        except UnicodeDecodeError:
            continue

    raise DocumentParserError(
        f"Could not decode text file '{filename}' with any supported encoding"
    )
